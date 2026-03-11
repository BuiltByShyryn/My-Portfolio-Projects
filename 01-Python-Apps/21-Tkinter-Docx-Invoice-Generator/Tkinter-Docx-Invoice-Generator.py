import tkinter as tk
from tkinter import messagebox
from tkinter.simpledialog import askstring
from tkinter import simpledialog
from docx import Document

root = tk.Tk()
root.title("Shop")

cart = []

def add_to_cart(product_name,price):
    cart.append({"name": product_name, "price": price})
    messagebox.showinfo("added", f"{product_name} was added in cart")

product_name = "Product 1"
price = 100

add_button=tk.Button(root, text="Add to a shopping cart", command=lambda:add_to_cart(product_name, price))
add_button.pack()

def view_cart():
    cart_window = tk.Toplevel(root)
    cart_window.title("cart")

    listbox = tk.Listbox(cart_window,height=10, width=50)
    listbox.pack()
    
    total_price = 0
    
    for item in cart:
        listbox.insert(tk.END, f"{item['name']} - {item['price']}")
        total_price +=item['price']
        
    listbox.insert(tk.END, f"the PRice of all item: {total_price}")
    
    def delete_item():
        try:
            chosen_item = listbox.curselection()[0]
            item_to_remove = cart[chosen_item]
            cart.remove(item_to_remove)
            listbox.delete(chosen_item)
            messagebox.showinfo("deeleted", f"{item_to_remove['name']} was deleted")
        except IndexError:
            messagebox.showwarning("warning", "choose the product to delete")
            
delete_button = tk.Button(root, text="Delete button", command = delete_item)
delete_button.pack()

def clear_cart():
    cart.clear()
    listbox.delete(0, tk.END)
    messagebox.showinfo("cleaned", "the cart is cleaned")

    clear_button = tk.Button(root, text="cleaning button", command = clear_cart)
    clear_button.pack()


    def generate():
        first_name = askstring("Clients Info", "type ur name ")
        last_name = askstring("Clients Info", "type ur surname")
        address = askstring("Clients Info" , "type ur address")
        phone = askstring("Clients Info", "Type ur phone number")
        email = askstring("Clients Info", "type uremail:")
    
        doc = Document()
        doc.add_heading("Nakladnaya", 0)
        doc.add_paragraph(f"Name: {first_name}")
        doc.add_paragraph(f"Surname: {last_name}")
        doc.add_paragraph(f"Address:{address}")
        doc.add_paragraph(f"Phone:{phone}")
        doc.add_paragraph(f"email: {email}") 
    
        doc.add_paragraph("products in cart")
        total_price = 0
        for item in cart:
            doc.add_paragraph(f"{item['name']} - {item['price']}")
            total_price+=item['price']
        
        doc.add_paragraph(f"Sum: {total_price}")
    
        doc.save("NAKLADNAYA.docx")
        messagebox.showinfo("nakladnaya", "Nakladnaya was saved")
    
    creating_button = tk.Button(root, text="Make a new Nakldanya", command=generate)
    creating_button.pack()

see_button = tk.Button(root, text="See teh cart", command=view_cart)
see_button.pack()

root.mainloop()
