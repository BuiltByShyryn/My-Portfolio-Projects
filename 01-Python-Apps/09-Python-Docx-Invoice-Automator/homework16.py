from docx import Document
import json



def load():
    with open("homework16/cart.json", "r") as file:
        data = json.load(file)
    
    return data

def create(orders):
    
    doc = Document()
    doc.add_heading("Накладная для всех клиентов", 0)
    for order in orders:
        doc.add_heading(f"Накладная для {order['client']['name']}", level=1)
        
        doc.add_paragraph(f"Имя клиента: {order['client']['name']}")
        doc.add_paragraph(f"Электронная почта: {order['client']['email']}")
        doc.add_paragraph(f"Номер заказа: {order['order_id']}")
        doc.add_paragraph(f"Дата заказа: {order['date']}")
    
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Товар'
        hdr_cells[1].text = 'Цена'
        hdr_cells[2].text = 'Количество'
        hdr_cells[3].text = 'Сумма'
        
        
        total_sum = 0
        
        for item in order['cart']:
            product = item['name']
            price = item['price']
            quantity = item['quantity']
            total = price * quantity
            total_sum += total
        
        
            row_cells = table.add_row().cells
            row_cells[0].text = product
            row_cells[1].text = f"{price} "
            row_cells[2].text = str(quantity)
            row_cells[3].text = f"{total} "
        
    doc.add_paragraph(f"\ntotal sum:{total_sum} ")
    
    

    doc.save("final.docx")
    print("file saved")
    
def generate():
    orders = load()  
    create(orders) 

generate()
